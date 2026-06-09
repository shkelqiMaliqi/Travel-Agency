from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Lenovo\OneDrive\Desktop\Travel Agency")
OUT = ROOT / "Travel_Agency_Raport_Teknik.docx"

BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF6FC"
GREEN = "E2F0D9"
YELLOW = "FFF2CC"
RED = "FCE4D6"
GRAY = "F2F2F2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C9D6")


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color="FFFFFF", size=9)
        set_cell_shading(hdr[i], BLUE)
        if widths:
            hdr[i].width = widths[i]
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=8.5)
            if widths:
                cells[i].width = widths[i]
            if index % 2 == 0:
                set_cell_shading(cells[i], "FAFCFE")
    document.add_paragraph()
    return table


def add_status_table(document, rows):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = ["Kerkesa", "Statusi", "Koment teknik"]
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color="FFFFFF", size=9)
        set_cell_shading(table.rows[0].cells[i], BLUE)
    for requirement, status, comment in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], requirement, bold=True, size=8.3)
        set_cell_text(cells[1], status, bold=True, size=8.3)
        set_cell_text(cells[2], comment, size=8.3)
        if status == "Po":
            fill = GREEN
        elif status == "Pjeserisht":
            fill = YELLOW
        else:
            fill = RED
        set_cell_shading(cells[1], fill)
    document.add_paragraph()


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Faqe ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for name, size, color in [
        ("Title", 24, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_cover(document):
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPORT TEKNIK")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Travel Agency - Web Services dhe Web API")
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph()
    document.add_paragraph()

    meta = [
        ("Lenda", "Zhvillimi i Ueb Sherbimeve dhe Ueb API-ve"),
        ("Projekti", "Travel Agency Portal"),
        ("Frontend", "React"),
        ("Backend", "ASP.NET Core Web API"),
        ("Databaza", "SQL Server"),
        ("Studenti", "[Shkruaj emrin dhe mbiemrin]"),
        ("Grupi", "[Shkruaj grupin]"),
        ("Data", date.today().strftime("%d.%m.%Y")),
    ]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for key, value in meta:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, bold=True, size=10)
        set_cell_text(cells[1], value, size=10)
        set_cell_shading(cells[0], PALE_BLUE)
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Raport i pergatitur per dorezim akademik ne Moodle.")
    run.italic = True
    run.font.size = Pt(10)
    document.add_page_break()


def add_toc(document):
    document.add_heading("Tabela e permbajtjes", level=1)
    entries = [
        "1. Permbledhje Ekzekutive",
        "2. Qellimi dhe objektivat e projektit",
        "3. Analiza e kerkesave funksionale dhe jofunksionale",
        "4. Projektimi i sistemit",
        "5. Pershkrimi i implementimit",
        "6. Siguria dhe arsyetimet teknike",
        "7. Dokumentimi, versionimi dhe infrastruktura API",
        "8. Testimi dhe rezultatet",
        "9. DevOps, dorezimi dhe kufizimet",
        "10. Perfundime dhe rekomandime",
        "11. Referencat",
        "12. Shtojcat dhe deklarata e origjinalitetit",
    ]
    add_numbered(document, entries)
    document.add_page_break()


def build_report():
    document = Document()
    configure_styles(document)
    add_cover(document)

    section = document.sections[0]
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    add_toc(document)

    document.add_heading("1. Permbledhje Ekzekutive", level=1)
    document.add_paragraph(
        "Travel Agency Portal eshte nje sistem web per menaxhimin e destinacioneve turistike, hoteleve, paketave "
        "te udhetimit dhe rezervimeve. Projekti eshte zhvilluar si React frontend dhe ASP.NET Core Web API backend. "
        "Backend-i ekspozon REST API te versionuara me /api/v1, dokumentohet me Swagger/OpenAPI dhe perdor SQL Server "
        "per ruajtjen e te dhenave."
    )
    document.add_paragraph(
        "Sistemi perfshin autentifikim me JWT, autorizim sipas roleve admin/user, MFA per role te mbrojtura me queue/email delivery, "
        "validim te input-eve, password hashing, rate limiting, audit logging, distributed caching, RabbitMQ, Swagger examples, "
        "Prometheus metrics, ELK log shipping, Kubernetes manifests, Playwright E2E tests, Dockerfile dhe docker-compose baze/enterprise. "
        "Ai permbush kerkesat kryesore te lendes Web Services dhe Web API dhe demonstron ne menyre konkrete kerkesat e avancuara."
    )

    document.add_heading("2. Qellimi dhe objektivat e projektit", level=1)
    document.add_paragraph(
        "Qellimi i projektit eshte krijimi i nje platforme funksionale per agjenci udhetimi, ku perdoruesit mund te shohin "
        "destinacione dhe paketa, te bejne rezervime dhe te komunikojne me stafin, ndersa administratori menaxhon katalogun "
        "dhe te dhenat operative."
    )
    add_bullets(document, [
        "Te implementohet REST API me endpoint-e te qarta per users, places, hotels, packages, bookings dhe contact messages.",
        "Te perdoret JWT per autentifikim stateless dhe RBAC per ndarjen e lejeve admin/user.",
        "Te shtohet MFA per role te ndjeshme si admin me delivery permes RabbitMQ/email.",
        "Te perdoret SQL Server per model relacional te te dhenave.",
        "Te dokumentohet API me Swagger/OpenAPI dhe te perdoret versionim /api/v1.",
        "Te ofrohet React frontend qe konsumon API-ne per funksionalitetet e perdoruesit dhe administratorit.",
        "Te perfshihen praktika sigurie si input validation, parameterized SQL, password hashing dhe rate limiting.",
        "Te demonstrohen caching, monitoring, logging qendror, microservice, queue processing, Kubernetes dhe integrime te jashtme.",
    ])

    document.add_heading("3. Analiza e kerkesave funksionale dhe jofunksionale", level=1)
    document.add_heading("3.1 Kerkesat funksionale", level=2)
    add_bullets(document, [
        "Regjistrim dhe login i perdoruesve.",
        "Forgot password me reset code dhe vendosje te password-it te ri.",
        "Ndryshim password-i nga profili i user-it me verifikim te current password.",
        "Shfaqje dhe kerkim i destinacioneve dhe paketave.",
        "Admin mund te menaxhoje destinations, hotels dhe packages.",
        "User mund te beje booking per paketa dhe te shikoje rezervimet e veta.",
        "Admin mund te shikoje bookings, users dhe contact messages.",
        "Contact form dergon mesazhe te lexueshme nga admin dashboard.",
    ])
    document.add_heading("3.2 Kerkesat jofunksionale", level=2)
    add_bullets(document, [
        "Siguri permes JWT, RBAC, password hashing dhe validimit te te dhenave.",
        "Dokumentim interaktiv permes Swagger UI me shembuj kerkesash/pergjigjesh.",
        "Versionim i API-se permes /api/v1.",
        "Logging, audit logging dhe global exception handling per diagnostikim.",
        "Caching i shperndare me Redis-ready configuration.",
        "Asynchronous processing permes RabbitMQ dhe notification microservice.",
        "Monitoring me metrics endpoint, snapshots dhe konfigurim Prometheus/Grafana.",
        "Docker, Docker Compose dhe Kubernetes manifests per deployment.",
        "Playwright E2E tests per browser-level validation.",
        "Teste automatike bazike per health endpoint, authentication protection dhe password hashing.",
    ])

    document.add_heading("4. Projektimi i sistemit", level=1)
    document.add_heading("4.1 Arkitektura", level=2)
    document.add_paragraph(
        "Projekti eshte implementuar si monolithic REST API me ndarje modulare ne controllers, models, services, middleware "
        "dhe filters. Kjo qasje eshte e pershtatshme per projekt studentor sepse ul kompleksitetin e deployment-it dhe "
        "testimit, por ruan strukture te qarte per zgjerim te mevonshem."
    )
    add_table(document, ["Komponenti", "Teknologjia", "Roli"], [
        ("Frontend", "React", "Nderfaqe per user dhe admin; konsumon REST API."),
        ("Backend", "ASP.NET Core Web API", "Ekspozon endpoint-et REST dhe logjiken e biznesit."),
        ("Database", "SQL Server", "Ruajtja e users, places, hotels, packages, bookings dhe contact messages."),
        ("Documentation", "Swagger/OpenAPI", "Dokumentim dhe testim interaktiv i endpoint-eve."),
        ("Security", "JWT + RBAC", "Autentifikim dhe autorizim sipas roleve."),
        ("Observability", "Prometheus + audit logs", "Metrics dhe gjurme auditimi per monitorim."),
        ("Messaging", "RabbitMQ", "Queue per MFA delivery dhe async processing."),
        ("Microservice", "NotificationService", "Sherbim i ndare per dergim MFA me email."),
        ("Infrastructure", "NGINX + Redis + ELK + MinIO + Mongo + Kubernetes", "Gateway, cache, logging, deployment dhe integrime te jashtme."),
        ("Deployment", "Docker + Docker Compose", "Mjedis i izoluar per API dhe stack-un mbeshtetes."),
    ], widths=[Inches(1.4), Inches(1.7), Inches(3.3)])

    document.add_heading("4.2 Modeli i te dhenave", level=2)
    document.add_paragraph(
        "Databaza eshte relacionale. Paketat lidhen me destinacione dhe hotele, ndersa bookings lidhen me user dhe package. "
        "Contact_Form ruan mesazhet e derguara nga vizitoret ose user-at e loguar."
    )
    add_table(document, ["Tabela", "Pershkrimi"], [
        ("Users", "Te dhenat e user-ave, rolet, username, email dhe password hash."),
        ("Places", "Destinacionet turistike."),
        ("Hotels", "Hotelet e lidhura me destinacione."),
        ("Travel_Packages", "Paketat e udhetimit me hotel, destinacion, cmim, data dhe vende te lira."),
        ("Bookings", "Rezervimet e user-ave per paketa."),
        ("Contact_Form", "Mesazhet e kontaktit qe lexohen nga admin."),
        ("Password_Reset_Codes", "Kodet e perkohshme per reset password."),
        ("User_Mfa_Codes", "Kodet e perkohshme per MFA challenge."),
        ("Audit_Logs", "Ngjarjet e audituara te kerkesave dhe autentifikimit."),
        ("Metrics_Snapshots", "Statistika periodike te ruajtura nga background worker."),
    ], widths=[Inches(1.8), Inches(4.6)])

    document.add_heading("5. Pershkrimi i implementimit", level=1)
    document.add_paragraph(
        "Backend-i perdor ASP.NET Core controllers per ekspozimin e endpoint-eve. API kthen pergjigje te unifikuara me "
        "fushat success, message dhe data. Global exception middleware kthen gabimet ne format JSON te lexueshem, ndersa "
        "request logging middleware regjistron metoden, path-in, status code dhe kohezgjatjen."
    )
    add_table(document, ["Endpoint", "Autorizimi", "Qellimi"], [
        ("POST /api/v1/auth/register", "Public", "Regjistrim i user-it."),
        ("POST /api/v1/auth/login", "Public", "Login dhe kthim JWT ose MFA challenge me queue/email delivery."),
        ("POST /api/v1/auth/verify-mfa", "Public", "Verifikim MFA dhe leshim i JWT final."),
        ("POST /api/v1/auth/forgot-password", "Public", "Gjenerim reset code."),
        ("POST /api/v1/auth/reset-password", "Public", "Ndryshim password-i me reset code."),
        ("GET /api/v1/packages", "Public", "Lista/filter i paketave."),
        ("POST /api/v1/bookings", "User", "Krijim booking per package."),
        ("GET /api/v1/bookings", "Admin", "Shfaqja e te gjitha bookings."),
        ("GET /api/v1/users", "Admin", "Lista e perdoruesve."),
        ("PUT /api/v1/users/{id}/password", "User", "Ndryshim password-i nga profili."),
        ("GET /api/v1/infrastructure/audit-logs", "Admin", "Shfaq audit logs te ruajtura."),
        ("POST /api/v1/infrastructure/storage-demo", "Admin", "Demonstrim i ruajtjes ne S3-compatible storage."),
    ], widths=[Inches(2.6), Inches(1.2), Inches(2.6)])

    document.add_heading("6. Siguria dhe arsyetimet teknike", level=1)
    document.add_heading("6.1 Pse REST dhe jo SOAP", level=2)
    document.add_paragraph(
        "REST u zgjodh sepse eshte me i pershtatshem per aplikacione moderne web me React frontend. REST perdor HTTP methods "
        "si GET, POST, PUT dhe DELETE, dhe kthen te dhena ne JSON. JSON eshte i lehte per JavaScript/React dhe per testim ne Swagger."
    )
    document.add_paragraph(
        "SOAP perdor XML, WSDL dhe struktura me te renda. Ai eshte i pershtatshem per sisteme enterprise/legacy me kontrata "
        "shume strikte, por per Travel Agency me resources si users, hotels, packages dhe bookings, REST eshte me praktik dhe me i thjeshte."
    )

    document.add_heading("6.2 Pse JWT", level=2)
    document.add_paragraph(
        "JWT u zgjodh sepse REST API duhet te jete stateless. Pas login-it, backend-i gjeneron token dhe frontend-i e dergon ate "
        "ne header-in Authorization: Bearer <token>. Serveri nuk ka nevoje te ruaje session per cdo user."
    )
    add_bullets(document, [
        "JWT punon mire me REST API dhe React frontend.",
        "API mund te lexoje user id dhe role nga token-i.",
        "Lejon ndarjen e autorizimeve admin/user.",
        "E ben sistemin me te pershtatshem per shkallezim ne te ardhmen.",
    ])

    document.add_heading("6.3 RBAC dhe mbrojtja e te dhenave", level=2)
    document.add_paragraph(
        "Role-Based Access Control ndan funksionet e user-it normal nga funksionet administrative. Admin menaxhon katalogun, "
        "bookings, users dhe messages. User mund te beje booking, te shikoje rezervimet e veta, te ndryshoje profilin/password "
        "dhe te dergoje mesazhe kontakti."
    )
    add_bullets(document, [
        "Password-et ruhen me salted PBKDF2 hashing.",
        "SQL queries perdorin parametra per te ulur rrezikun e SQL Injection.",
        "Input validation aplikohet ne request models.",
        "Rate limiting aplikohet per login, password reset dhe public writes.",
        "MFA aplikohet per role te konfiguruara si admin dhe mund te dergohet permes RabbitMQ/email.",
        "Audit logs ruhen ne databaze per qellime gjurmimi.",
        "Contact form lidh user-in nga JWT token, jo nga U_Id i derguar nga browser-i.",
    ])

    document.add_heading("7. Dokumentimi, versionimi dhe infrastruktura API", level=1)
    document.add_paragraph(
        "API eshte e dokumentuar me Swagger/OpenAPI dhe mund te testohet nga /swagger. Dokumentimi tani perfshin bearer auth, "
        "shembuj kerkesash/pergjigjesh dhe kode gabimesh te zakonshme. Endpoint-et jane versionuar ne URL me /api/v1, "
        "qe lejon shtimin e /api/v2 ne te ardhmen pa prishur klientet ekzistues."
    )
    add_bullets(document, [
        "Swagger UI: http://localhost:5132/swagger",
        "Health check: http://localhost:5132/health",
        "Metrics: http://localhost:5132/metrics",
        "Frontend: http://localhost:3000",
        "API base URL: http://localhost:5132/api/v1",
    ])

    document.add_heading("8. Testimi dhe rezultatet", level=1)
    document.add_paragraph(
        "Testimi eshte bere ne dy menyra: teste automatike backend dhe testim manual permes Swagger/frontend. Testet automatike "
        "kontrollojne health endpoint, mbrojtjen e endpoint-eve te siguruara dhe password hashing."
    )
    add_table(document, ["Testi", "Lloji", "Rezultati i pritur"], [
        ("HealthEndpoint_ReturnsHealthy", "Automatik", "GET /health kthen 200 OK."),
        ("ProtectedEndpoint_RequiresAuthentication", "Automatik", "GET /api/v1/users pa token kthen 401 Unauthorized."),
        ("PasswordHasherTests", "Automatik", "Hash i ri verifikohet dhe password i gabuar refuzohet."),
        ("Swagger example flow", "Manual", "Login/MFA demo testohet direkt ne Swagger."),
        ("Playwright E2E", "Automatik", "Browser teston home, login dhe forgot-password flow."),
        ("Register/Login", "Manual", "User krijohet dhe merr JWT token."),
        ("Booking flow", "Manual", "User ben booking dhe vendet e lira ulen."),
        ("Infrastructure demo", "Manual", "Audit logs, metrics snapshots dhe storage endpoint demonstrohen."),
    ], widths=[Inches(2.3), Inches(1.2), Inches(2.9)])

    document.add_heading("9. DevOps, dorezimi dhe kufizimet", level=1)
    document.add_paragraph(
        "Projekti perfshin Dockerfile, docker-compose.yml, docker-compose.enterprise.yml dhe Kubernetes manifests per ekzekutim baze "
        "dhe per demonstrim te gateway/load balancing, Redis, RabbitMQ, notification microservice, MongoDB, MinIO, ELK, Prometheus, "
        "Grafana, Alertmanager dhe automated backups. Gjithashtu ekzistojne workflow GitHub Actions per backend, frontend dhe E2E CI."
    )
    document.add_heading("9.1 Statusi kundrejt kerkesave teknike", level=2)
    add_status_table(document, [
        ("REST API", "Po", "Endpoint-et perdorin HTTP methods dhe JSON."),
        ("Swagger/OpenAPI", "Po", "Dokumentim interaktiv ne /swagger me shembuj dhe bearer auth."),
        ("JWT Authentication", "Po", "Login gjeneron JWT dhe endpoint-et e mbrojtura kerkojne Bearer token."),
        ("RBAC", "Po", "Role admin/user per autorizim."),
        ("MFA", "Po", "MFA per role te konfiguruara, me queue/email delivery dhe local demo mode."),
        ("SQL Server", "Po", "Databaze relacionale me tables per user, packages, bookings, contact."),
        ("ORM", "Po", "EF Core perdoret per support tables si audit logs dhe MFA."),
        ("Docker", "Po", "Dockerfile dhe docker-compose.yml ekzistojne."),
        ("CI/CD", "Po", "Ekzistojne workflow backend dhe frontend per build/test."),
        ("Redis/Advanced cache", "Po", "Distributed cache me Redis-ready configuration dhe fallback local."),
        ("Load balancing", "Po", "NGINX gateway me dy API instances ne compose enterprise."),
        ("Prometheus/Grafana", "Po", "Metrics endpoint dhe konfigurim monitorimi ekzistojne."),
        ("ELK Logging", "Po", "API mund te dergoje JSON logs ne Logstash permes TCP."),
        ("Asynchronous processing", "Po", "RabbitMQ dhe NotificationService procesojne MFA delivery ne sfond."),
        ("API Gateway", "Po", "NGINX gateway i perfshire ne stack."),
        ("Microservices", "Po", "NotificationService eshte microservice i ndare nga API kryesore."),
        ("Kubernetes", "Po", "Manifests per deployments, services, ingress, HPA dhe backup CronJob."),
        ("Automated backups", "Po", "Docker backup helper dhe Kubernetes CronJob ekzistojne."),
        ("E2E tests", "Po", "Playwright E2E tests dhe CI workflow ekzistojne."),
    ])

    document.add_heading("10. Perfundime dhe rekomandime", level=1)
    document.add_paragraph(
        "Travel Agency Portal permbush qellimin kryesor te projektit: ofron REST API funksionale, te dokumentuar dhe te siguruar "
        "per menaxhimin e nje agjencie udhetimi. Projekti demonstron perdorimin e Web API, JWT, RBAC, MFA, SQL Server, Swagger, "
        "Docker, RabbitMQ, microservices, Kubernetes, caching, monitoring dhe React frontend qe konsumon sherbimet."
    )
    document.add_paragraph(
        "Per nje version production/enterprise rekomandohet konfigurimi i sekreteve reale, TLS certificates, SMTP/webhook providers, "
        "cloud deployment dhe restore drills periodike."
    )

    document.add_heading("11. Referencat", level=1)
    add_bullets(document, [
        "Microsoft Docs - ASP.NET Core Web API.",
        "Microsoft Docs - JWT Bearer Authentication.",
        "OpenAPI Specification dhe Swagger UI documentation.",
        "Microsoft SQL Server documentation.",
        "Docker documentation.",
        "React documentation.",
    ])

    document.add_heading("12. Shtojcat dhe deklarata e origjinalitetit", level=1)
    document.add_heading("12.1 Shtojca A - Udhezime ekzekutimi", level=2)
    add_bullets(document, [
        "Backend: cd \"Travel Agency Portal/Travel Agency Portal\" dhe dotnet run --launch-profile http.",
        "Frontend: cd \"UI/travel-agency\" dhe npm start.",
        "Swagger: http://localhost:5132/swagger.",
        "Admin login: admin@travelagency.com / Admin123!.",
        "Admin MFA: ne local demo kodi mund te kthehet ne pergjigje; ne production dergohet permes RabbitMQ/email.",
        "Kubernetes: kubectl apply -k Infrastructure/kubernetes.",
        "E2E: cd UI/travel-agency dhe npm run e2e.",
    ])
    document.add_heading("12.2 Deklarata e origjinalitetit", level=2)
    document.add_paragraph(
        "Une, [Shkruaj emrin dhe mbiemrin], deklaroj se ky projekt eshte punuar nga une/grupi im dhe se materiali i dorezuar "
        "nuk permban plagjiature. Burimet e perdorura jane referuar ne seksionin e referencave dhe kodi burimor eshte zhvilluar "
        "per qellime akademike ne kuader te lendes Zhvillimi i Ueb Sherbimeve dhe Ueb API-ve."
    )
    document.add_paragraph("Nenshkrimi: ____________________________")
    document.add_paragraph("Data: ____________________________")

    for section in document.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    document.core_properties.title = "Travel Agency - Raport Teknik"
    document.core_properties.subject = "Web Services dhe Web API"
    document.core_properties.author = "Travel Agency Project"
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_report()
    print(path)
